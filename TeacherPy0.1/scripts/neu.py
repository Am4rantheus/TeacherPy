
import os
# -*- coding: utf-8 -*-
import sys
import shutil
import json
import win32com.client
import pywintypes
import subprocess
from docx import Document
import time
from colorama import init, Fore, Style

# Füge den scripts-Ordner zum Python-Pfad hinzu
SCRIPT_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(SCRIPT_BASE_DIR)
sys.path.append(SCRIPT_BASE_DIR)

# Importiere die Pfade aus paths.py
from paths import SVP_SCRIPT_PATH, TEMPLATE_PATH, TEMPLATE_NOTE_PATH

init(autoreset=True)  # Initialize colorama

def load_config(file_name):
    json_path = os.path.join(PARENT_DIR, file_name)
    
    if not os.path.exists(json_path):
        print(f"Fehler: Die Konfigurationsdatei '{json_path}' wurde nicht gefunden.")
        exit(1)

    try:
        with open(json_path, 'r') as config_file:
            return json.load(config_file)
    except json.JSONDecodeError:
        print(f"Fehler: Die Datei '{json_path}' enthält kein gültiges JSON-Format.")
        exit(1)

config = load_config('project_config.json')
layout = load_config('layout.json')

def get_color(key):
    return getattr(Fore, layout['colors'].get(key, 'WHITE'))

def colored_filename(filename):
    if filename.startswith("SVP_"):
        return f"{get_color('SVP_')}{filename}{Style.RESET_ALL}"
    elif filename.endswith((".pptx", ".odp")):
        return f"{get_color('PRESENTATION')}{filename}{Style.RESET_ALL}"
    elif filename.endswith((".docx", ".odt")):
        return f"{get_color('DOCX')}{filename}{Style.RESET_ALL}"
    elif filename.endswith(".pdf"):
        if "_Erwartungsbild" in filename:
            return f"{get_color('ERWARTUNGSBILD')}{filename}{Style.RESET_ALL}"
        else:
            return f"{get_color('PDF')}{filename}{Style.RESET_ALL}"
    elif filename.endswith((".mp4", ".avi", ".mov", ".wmv")):
        return f"{get_color('VIDEO')}{filename}{Style.RESET_ALL}"
    else:
        for prefix in ['AB_', 'LB_']:
            if filename.startswith(prefix):
                return f"{get_color(prefix)}{filename}{Style.RESET_ALL}"
    return filename

def colored_foldername(foldername):
    return f"{Fore.YELLOW}{foldername}{Style.RESET_ALL}"

def toggle_abbreviations(doc_path, include):
    doc = Document(doc_path)
    
    abbreviations = "verwendete Abkürzungen:\nE/M: Einstieg/Motiviation , V: Vermittlung, VT: Vertiefung, Erar: Erarbeitung, ErgS: Ergebnissicherung, DidRes: Didaktische Reserve StE: Stundenergebnis , UG: Unterrichtsgespräch, EZ: Einzelarbeit, PA: Partnerarbeit, T: Tafelbild, LV: Lehrervortrag, LB: Lehrbuch"
    
    # Find the paragraph with abbreviations
    abbr_paragraphs = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("verwendete Abkürzungen:") or (abbr_paragraphs and paragraph.text.strip()):
            abbr_paragraphs.append(i)
        elif abbr_paragraphs:
            break

    if include:
        # If abbreviations should be included
        if abbr_paragraphs:
            # Update existing abbreviations
            doc.paragraphs[abbr_paragraphs[0]].text = abbreviations
            # Remove any additional paragraphs
            for i in reversed(abbr_paragraphs[1:]):
                doc._element.body.remove(doc.paragraphs[i]._element)
        else:
            # Add abbreviations if they don't exist
            doc.add_paragraph(abbreviations)
    else:
        # If abbreviations should not be included, remove all related paragraphs
        for i in reversed(abbr_paragraphs):
            doc._element.body.remove(doc.paragraphs[i]._element)
    
    doc.save(doc_path)

def create_project_structure(config, layout):
    # Ask for project name
    project_name = input("Bitte vergeben Sie einen Namen für die neue Stunde: ")

    # Define paths from config and paths.py
    base_folder_path = config['base_folder_path']
    notes_folder_path = config['notes_folder_path']

    # Create subfolder in base folder
    project_folder_path = os.path.join(base_folder_path, project_name)
    os.makedirs(project_folder_path, exist_ok=True)
    print(f"Unterordner {colored_foldername(project_name)} wurde im Ordner {colored_foldername(base_folder_path)} erstellt.")

    # Create Ressources subfolder in the project folder
    resources_folder_path = os.path.join(project_folder_path, "Ressources")
    os.makedirs(resources_folder_path, exist_ok=True)
    print(f"Unterordner {colored_foldername('Ressources')} wurde im Projektordner {colored_foldername(project_name)} erstellt.")

    # Convert .dotx to .docx and save in project folder
    word_file_name = f"SVP_{project_name}.docx"
    word_file_path = os.path.join(project_folder_path, word_file_name)
    
    word = win32com.client.Dispatch("Word.Application")
    try:
        doc = word.Documents.Add(TEMPLATE_PATH)
        doc.SaveAs(word_file_path, FileFormat=16)  # FileFormat=16 is for .docx
        doc.Close()
    except pywintypes.com_error as e:
        print(f"COM-Fehler aufgetreten: {e}")
        print(f"TEMPLATE_PATH: {TEMPLATE_PATH}")
        print(f"word_file_path: {word_file_path}")
        # Hier könnten Sie zusätzliche Fehlerbehandlung oder Logging hinzufügen
    finally:
        word.Quit()
    
    print(f"Word-Datei {colored_filename(word_file_name)} wurde im Ordner {colored_foldername(project_folder_path)} erstellt.")

    # Wait for the file to be accessible
    max_attempts = 10
    attempts = 0
    while attempts < max_attempts:
        if os.path.exists(word_file_path) and not os.path.isdir(word_file_path):
            try:
                with open(word_file_path, 'rb'):  # Try to open the file
                    break  # If successful, break the loop
            except PermissionError:
                pass  # If file is still locked, continue waiting
        time.sleep(1)  # Wait for 1 second before trying again
        attempts += 1

    if attempts == max_attempts:
        print(f"Fehler: Konnte nicht auf die Datei {colored_filename(word_file_name)} zugreifen. Bitte überprüfen Sie, ob die Datei existiert und nicht von einem anderen Programm verwendet wird.")
        return

    # Toggle abbreviations based on config
    toggle_abbreviations(word_file_path, config['svp_options']['include_abbreviations'])

    # Check if note creation is enabled in config
    if config['create_note'] and notes_folder_path:
        # Create .note file in notes folder
        note_file_name = f"Note_{project_name}.note"
        note_file_path = os.path.join(notes_folder_path, note_file_name)
        shutil.copy(TEMPLATE_NOTE_PATH, note_file_path)
        print(f".note-Datei {colored_filename(note_file_name)} wurde im Ordner {colored_foldername(notes_folder_path)} erstellt.")

    # Execute SVP.py
    subprocess.run([sys.executable, SVP_SCRIPT_PATH, project_folder_path])

def main():
    config = load_config('project_config.json')
    layout = load_config('layout.json')
    create_project_structure(config, layout)

if __name__ == "__main__":
    main()
