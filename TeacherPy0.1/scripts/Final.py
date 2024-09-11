import os
import sys
import glob
import re
from PyPDF2 import PdfMerger
from docx import Document
import win32com.client
import json
import shutil
from colorama import init, Fore, Style

# Füge den scripts-Ordner zum Python-Pfad hinzu
SCRIPT_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(SCRIPT_BASE_DIR)
sys.path.append(SCRIPT_BASE_DIR)

# Initialize colorama
init(autoreset=True)

def load_config(file_name):
    config_path = os.path.join(PARENT_DIR, file_name)

    try:
        with open(config_path, 'r') as config_file:
            return json.load(config_file)
    except FileNotFoundError:
        print(f"Fehler: Die Konfigurationsdatei wurde nicht gefunden: {config_path}")
        print(f"Bitte stellen Sie sicher, dass die '{file_name}'-Datei im übergeordneten Verzeichnis des Skripts liegt.")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Fehler: Die Datei '{config_path}' enthält kein gültiges JSON-Format.")
        sys.exit(1)

config = load_config('project_config.json')
layout = load_config('layout.json')
BASE_DIR = config['base_folder_path']
INBOX_DIR = config['notes_folder_path']
USB_PATH = config['usb_path']

def get_color(key):
    return getattr(Fore, layout['colors'].get(key, 'WHITE'))

def colored_filename(filename):
    if filename.startswith("SVP_"):
        return f"{get_color('SVP_')}{filename}{Style.RESET_ALL}"
    elif filename.startswith("AB_"):
        if "_Erwartungsbild" in filename:
            return f"{get_color('AB_LIGHT')}{filename}{Style.RESET_ALL}"
        return f"{get_color('AB_')}{filename}{Style.RESET_ALL}"
    elif filename.startswith("LB_"):
        if "_Erwartungsbild" in filename:
            return f"{get_color('LB_LIGHT')}{filename}{Style.RESET_ALL}"
        return f"{get_color('LB_')}{filename}{Style.RESET_ALL}"
    elif filename.endswith((".pptx", ".odp")):
        return f"{get_color('PRESENTATION')}{filename}{Style.RESET_ALL}"
    elif filename.endswith(".pdf"):
        return f"{get_color('PDF')}{filename}{Style.RESET_ALL}"
    return filename

def colored_foldername(foldername):
    return f"{get_color('FOLDER')}{foldername}{Style.RESET_ALL}"

def get_subfolder():
    subfolders = [f for f in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, f))]
    print("Verfügbare Unterordner:")
    for i, folder in enumerate(subfolders, 1):
        print(f"{i}. {colored_foldername(folder)}")
    while True:
        try:
            choice = int(input("Wählen Sie einen Unterordner (Nummer): ")) - 1
            if 0 <= choice < len(subfolders):
                return os.path.join(BASE_DIR, subfolders[choice])
            else:
                print("Ungültige Auswahl. Bitte versuchen Sie es erneut.")
        except ValueError:
            print("Bitte geben Sie eine Zahl ein.")

def set_default_font(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'

def convert_to_pdf(folder):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    
    docx_files = glob.glob(os.path.join(folder, "*.docx"))
    converted_files = []
    for docx in docx_files:
        filename = os.path.basename(docx)
        if (filename.startswith(("SVP_", "LB_", "AB_")) and 
            not filename.endswith("_old.docx")):
            pdf_name = os.path.splitext(docx)[0] + ".pdf"
            
            # Only ask for Erwartungsbild for AB_ and LB_ files
            include_comments = False
            if filename.startswith(("AB_", "LB_")):
                include_comments = input(f"Möchten Sie ein Erwartungsbild mit Kommentaren für {colored_filename(filename)} erstellen? (j/n): ").lower() == 'j'
            
            try:
                doc = Document(docx)
                set_default_font(doc)
                doc.save(docx)
                
                doc = word.Documents.Open(docx)
                
                # Version ohne Kommentare
                doc.ActiveWindow.View.ShowRevisionsAndComments = False
                doc.SaveAs(pdf_name, FileFormat=17)
                converted_files.append(os.path.basename(pdf_name))
                print(f"Konvertiert (ohne Kommentare): {colored_filename(filename)} -> {colored_filename(os.path.basename(pdf_name))}")
                
                # Version mit Kommentaren, wenn gewünscht
                if include_comments:
                    doc.ActiveWindow.View.ShowRevisionsAndComments = True
                    comment_pdf_name = os.path.splitext(docx)[0] + "_Erwartungsbild.pdf"
                    doc.SaveAs(comment_pdf_name, FileFormat=17)
                    converted_files.append(os.path.basename(comment_pdf_name))
                    print(f"Konvertiert (mit Kommentaren): {colored_filename(filename)} -> {colored_filename(os.path.basename(comment_pdf_name))}")
                
                doc.Close()
            except Exception as e:
                print(f"Fehler bei der Konvertierung von {colored_filename(filename)}: {e}")
  
    word.Quit()
    print("Konvertierung abgeschlossen.")
    print("Konvertierte Dateien:", ", ".join(colored_filename(f) for f in converted_files))
    return converted_files

def should_merge():
    while True:
        response = input("Möchten Sie die konvertierten Dateien zusammenführen? (j/n): ").lower()
        if response in ['j', 'n']:
            return response == 'j'
        print("Bitte antworten Sie mit 'j' oder 'n'.")

def get_order_from_docx(svp_file, converted_files):
    doc = Document(svp_file)
    order = []
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for filename in converted_files:
                    base_filename = filename.replace("_Erwartungsbild.pdf", "").replace(".pdf", "")
                    if base_filename in cell.text:
                        if base_filename not in order:
                            order.append(base_filename)
    
    print(f"Gefundene Dateinamen in der SVP-Datei: {', '.join(colored_filename(f) for f in order)}")
    return order

def merge_pdfs(folder, converted_files):
    svp_file = next((f for f in converted_files if f.startswith("SVP_")), None)
    if not svp_file:
        print("Keine SVP_-Datei gefunden.")
        return

    svp_pdf_path = os.path.join(folder, svp_file)
    if not os.path.exists(svp_pdf_path):
        print(f"Fehler: Die Datei {colored_filename(svp_pdf_path)} wurde nicht gefunden.")
        return

    merger = PdfMerger()
    
    merger.append(svp_pdf_path)
    print(f"SVP-Datei hinzugefügt: {colored_filename(svp_file)}")

    order = get_order_from_docx(svp_pdf_path.replace(".pdf", ".docx"), converted_files)
    
    for file_prefix in order:
        erwartungsbild_file = next((f for f in converted_files if f.startswith(file_prefix) and f.endswith("_Erwartungsbild.pdf")), None)
        if erwartungsbild_file:
            matching_pdf_path = os.path.join(folder, erwartungsbild_file)
            merger.append(matching_pdf_path)
            print(f"Hinzugefügt (mit Kommentaren): {colored_filename(erwartungsbild_file)}")
        else:
            matching_file = next((f for f in converted_files if f.startswith(file_prefix) and f.endswith(".pdf") and not f.endswith("_Erwartungsbild.pdf")), None)
            if matching_file:
                matching_pdf_path = os.path.join(folder, matching_file)
                merger.append(matching_pdf_path)
                print(f"Hinzugefügt: {colored_filename(matching_file)}")
            else:
                print(f"Warnung: Keine passende PDF für '{colored_filename(file_prefix)}' gefunden.")

    output_name = os.path.basename(svp_file)
    output_path = os.path.join(folder, output_name)
    
    if merger.pages:
        merger.write(output_path)
        print(f"PDFs wurden in {colored_filename(output_name)} zusammengeführt.")
    else:
        print("Keine PDFs zum Zusammenführen gefunden.")
    
    merger.close()
    return output_path

def copy_to_inbox(file_name):
    while True:
        response = input(f"Möchten Sie die finale Datei '{colored_filename(os.path.basename(file_name))}' in den INBOX-Ordner kopieren? (j/n): ").lower()
        if response == 'j':
            shutil.copy(file_name, INBOX_DIR)
            print(f"Datei wurde nach {colored_foldername(INBOX_DIR)} kopiert.")
            break
        elif response == 'n':
            print("Datei wurde nicht kopiert.")
            break
        else:
            print("Bitte antworten Sie mit 'j' oder 'n'.")

def copy_to_usb(folder):
    while True:
        print("\nMöchten Sie Dateien auf den USB-Stick kopieren?")
        print("1. Gesamten Ordner")
        print("2. Nur PDF-Dateien und Präsentationen")
        print("3. Nichts kopieren")
        choice = input("Wählen Sie eine Option (1/2/3): ")

        folder_name = os.path.basename(folder)
        usb_folder = os.path.join(USB_PATH, folder_name)

        if choice == '1':
            shutil.copytree(folder, usb_folder, dirs_exist_ok=True)
            print(f"Gesamter Ordner wurde auf den USB-Stick kopiert: {colored_foldername(usb_folder)}")
            break
        elif choice == '2':
            os.makedirs(usb_folder, exist_ok=True)
            for file in os.listdir(folder):
                if file.endswith(('.pdf', '.pptx')):
                    shutil.copy(os.path.join(folder, file), usb_folder)
            print(f"PDF-Dateien und Präsentationen wurden auf den USB-Stick kopiert: {colored_foldername(usb_folder)}")
            break
        elif choice == '3':
            print("Es wurden keine Dateien auf den USB-Stick kopiert.")
            break
        else:
            print("Ungültige Auswahl. Bitte wählen Sie 1, 2 oder 3.")

def main():
    folder = get_subfolder()
    converted_files = convert_to_pdf(folder)
    if converted_files and should_merge():
        final_file = merge_pdfs(folder, converted_files)
        if final_file:
            copy_to_inbox(final_file)
    else:
        print("Keine Dateien zum Zusammenführen oder Zusammenführen nicht gewünscht.")
    
    copy_to_usb(folder)

if __name__ == "__main__":
    main()