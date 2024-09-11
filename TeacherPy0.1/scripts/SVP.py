import os
import sys
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import re
import json
from colorama import init, Fore, Style

# Füge den scripts-Ordner zum Python-Pfad hinzu
SCRIPT_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(SCRIPT_BASE_DIR)
sys.path.append(SCRIPT_BASE_DIR)

# Importiere die Pfade aus paths.py (falls benötigt)
# from paths import SOME_PATH

# Initialize colorama
init(autoreset=True)

def load_config(file_name):
    config_path = os.path.join(PARENT_DIR, file_name)

    try:
        with open(config_path, 'r') as config_file:
            return json.load(config_file)
    except FileNotFoundError:
        print(f"Fehler: Die Konfigurationsdatei wurde nicht gefunden: {config_path}")
        print(f"Bitte stellen Sie sicher, dass die '{file_name}'-Datei im übergeordneten Verzeichnis des Skripts liegt.")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Fehler: Die Datei '{config_path}' enthält kein gültiges JSON-Format.")
        sys.exit(1)

config = load_config('project_config.json')
layout = load_config('layout.json')

def get_color(key):
    return getattr(Fore, layout['colors'].get(key, 'WHITE'))

def colored_filename(filename):
    if filename.startswith("SVP_"):
        return f"{get_color('SVP_')}{filename}{Style.RESET_ALL}"
    elif filename.endswith((".docx", ".odt")):
        return f"{get_color('DOCX')}{filename}{Style.RESET_ALL}"
    return filename

def colored_foldername(foldername):
    return f"{get_color('FOLDER')}{foldername}{Style.RESET_ALL}"

def find_svp_file(directory):
    for file in os.listdir(directory):
        if file.startswith("SVP_") and file.endswith(".docx"):
            return os.path.join(directory, file)
    return None

def update_table(doc, updates):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in updates.items():
                    if key in cell.text:
                        pattern = f"{re.escape(key)}.*"
                        cell.text = re.sub(pattern, f"{key} {value}", cell.text, flags=re.DOTALL)
                        
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

def update_time_column(doc, start_time, end_time):
    for table in doc.tables:
        if table.cell(0, 0).text.strip().lower() == "zeit":
            column = table.columns[0]
            first_empty_cell = None
            last_cell = None
            
            for cell in column.cells[1:]:  # Skip header
                if cell.text.strip() == "" and first_empty_cell is None:
                    first_empty_cell = cell
                last_cell = cell
            
            if first_empty_cell:
                first_empty_cell.text = start_time
            if last_cell:
                last_cell.text = end_time
            
            # Set font size to 8
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

def get_time_slot(time_slots):
    print("Wählen Sie einen Zeitraum:")
    for key, value in time_slots.items():
        print(f"{key}. {value[0]}-{value[1]}")
    choice = input("Ihre Wahl (1-4): ")
    return time_slots.get(choice, ["", ""])

def toggle_abbreviations(doc, include):
    abbreviations_title = "verwendete Abkürzungen:"
    abbreviations_content = "E/M: Einstieg/Motiviation , V: Vermittlung, VT: Vertiefung, Erar: Erarbeitung, ErgS: Ergebnissicherung, DidRes: Didaktische Reserve StE: Stundenergebnis , UG: Unterrichtsgespräch, EZ: Einzelarbeit, PA: Partnerarbeit, T: Tafelbild, LV: Lehrervortrag, LB: Lehrbuch"
    
    # Find the paragraph with abbreviations
    abbr_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(abbreviations_title):
            abbr_paragraph = paragraph
            break
    
    if include:
        if abbr_paragraph:
            # Clear existing content
            abbr_paragraph.clear()
        else:
            # Add new paragraph if not exists
            abbr_paragraph = doc.add_paragraph()
        
        # Add title
        title_run = abbr_paragraph.add_run(abbreviations_title)
        title_run.bold = True
        title_run.font.size = Pt(7)
        
        # Add line break
        abbr_paragraph.add_run("\n")
        
        # Add content
        content_run = abbr_paragraph.add_run(abbreviations_content)
        content_run.font.size = Pt(7)
        
        # Set paragraph alignment
        abbr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:
        # If abbreviations should not be included, remove the paragraph
        if abbr_paragraph:
            p = abbr_paragraph._element
            p.getparent().remove(p)

def main(directory, config):
    file_path = find_svp_file(directory)
    
    if not file_path:
        print("Keine SVP_*.docx Datei gefunden.")
        return

    # Rename the original file
    old_file_path = file_path.replace(".docx", "_old.docx")
    os.rename(file_path, old_file_path)

    doc = Document(old_file_path)

    svp_options = config['svp_options']
    time_slots = svp_options['time_slots']
    prompts = svp_options['prompts']
    default_teacher = svp_options['default_teacher']
    include_abbreviations = svp_options['include_abbreviations']

    start_time, end_time = get_time_slot(time_slots)
    updates = {}

    if prompts['Klasse/Kurs']:
        updates["Klasse/Kurs:"] = input("Klasse/Kurs: ")
    if prompts['Raumnummer']:
        updates["Zeit/Raum:"] = f"{start_time}-{end_time} / {input('Raumnummer: ')}"
    if prompts['Datum']:
        updates["Datum:"] = input("Datum: ")
    if prompts['Lernbereich']:
        updates["Lernbereich:"] = input("Lernbereich: ")
    if prompts['Stundenthema']:
        updates["Stundenthema:"] = input("Stundenthema: ")
    if prompts['Lehrperson']:
        updates["Lehrperson:"] = input(f"Lehrperson [{default_teacher}]: ") or default_teacher

    update_table(doc, updates)
    update_time_column(doc, start_time, end_time)
    toggle_abbreviations(doc, include_abbreviations)

    # Save the updated document with the original name
    doc.save(file_path)
    print(f"Aktualisierte Datei gespeichert als: {colored_filename(os.path.basename(file_path))}")

    if os.name == 'nt':  # for Windows
        os.startfile(file_path)
    elif os.name == 'posix':  # for macOS and Linux
        subprocess.call(('open', file_path))

if __name__ == "__main__":
    if len(sys.argv) > 1:
        directory = sys.argv[1]
    else:
        directory = input("Geben Sie den Pfad zum Ausgangsordner ein: ")
    
    main(directory, config)